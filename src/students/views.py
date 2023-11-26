#from accounts.models import Shekh_Mohafez,Student_Apply
from django.shortcuts import render,redirect
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from .models import Shekh_Mohafez,Student_Apply


def student(request,stu_id):
    pass
def apply_compition(request):
    context = {
        'shekh_mohafez': Shekh_Mohafez.objects.all(),
    }
    if request.method== 'POST' and "btnapply" in request.POST:
        redirect(request,'apply_compition')
    else:
        return render(request,'students/apply_comptition.html',context)
def enter_results(request):
    return render(request,'students/enter_results.html')




def generate_word_response(students):
    document = Document()
    # Add a table to the document
    table = document.add_table(rows=1, cols=9)
    # Set the table alignment to right
    table.alignment = 2  # 2 corresponds to RIGHT alignment
    # Style the table
    table.style = 'TableGrid'
    # Add headers to the table
    headers = table.rows[0].cells
    for header in headers:
        # Set alignment for headers
        if not header.paragraphs:
            header.add_paragraph()
        if not header.paragraphs[0].runs:
            header.paragraphs[0].add_run()
        header.paragraphs[0].alignment = 2  # 2 corresponds to RIGHT alignment
        header.vertical_alignment = 0  # 0 corresponds to TOP alignment
        header.paragraphs[0].runs[0].font.size = Pt(16)  # Set font size for headers
        header.paragraphs[0].runs[0].font.bold = True  # Make headers bold
        # Set Arabic headers
        arabic_headers = ['النتيجة', 'العام', 'رقم التلفون الثاني', 'رقم التلفون الأول', 'عدد الأجزاء',
                          'اسم الشيخ المحفظ', 'البلد', 'الاسم', 'م']
        header.text = arabic_headers[headers.index(header)]
    # Set paragraph alignment for the entire table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 2  # 2 corresponds to RIGHT alignment
    # Add data to the table
    i = 0
    for student in students :
        i += 1
        row_cells = table.add_row().cells
        for cell in row_cells:
            # Create a new paragraph and run if not exist
            if not cell.paragraphs:
                cell.add_paragraph()
            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run()
            # Set alignment for data cells
            cell.paragraphs[0].alignment = 2  # 2 corresponds to RIGHT alignment
            cell.vertical_alignment = 0  # 0 corresponds to TOP alignment
            cell.paragraphs[0].runs[0].font.size = Pt(14)  # Set font size for data
        row_cells[8].text = str(i)
        row_cells[7].text = str(student.full_name)
        row_cells[6].text = str(student.student_country)
        row_cells[5].text = str(student.student_shekh_mohafez.user.first_name) + ' ' + str(
            student.student_shekh_mohafez.user.last_name)
        row_cells[4].text = str(student.number_of_Juze)
        row_cells[3].text = str(student.phone_number1)
        row_cells[2].text = str(student.phone_number2)
        row_cells[1].text = str(student.current_date)
        row_cells[0].text = str(student.result)
    # Adjust column widths based on content
    # Create a response with the Word document

    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = 'attachment; filename="exported_data.docx"'
    document.save(response)
    return response

def search(request):
    all_students = Student_Apply.objects.all()
    # Get the last stored values from the session
    original_name = request.session.get('original_name', None)
    original_noofguze = request.session.get('original_noofguze', None)
    original_id_shekh = request.session.get('original_id_shekh', None)
    original_cs = request.session.get('original_cs', 'off')
    # Apply filters based on user input
    cs='off'
    if 'cs' in request.GET:
        cs = request.GET['cs']
        if not cs:
            cs = 'off'
    if 'searchname' in request.GET:
        name = request.GET['searchname']
        if name:
            if cs == 'on':
                all_students = all_students.filter(full_name__contains=name)
            else:
                all_students = all_students.filter(full_name__icontains=name)
    if 'mohafez_select' in request.GET:
        id_shekh = request.GET['mohafez_select']
        if int(id_shekh) == 0:
            all_students = all_students
        else:
            all_students = all_students.filter(student_shekh_mohafez__id__icontains=id_shekh)
    if 'no_of_juze_search' in request.GET:
        noofguze = request.GET['no_of_juze_search']
        if int(noofguze) == 0:
            all_students = all_students
        else:
            all_students = all_students.filter(number_of_Juze__icontains=noofguze)
    context = {
        'shekh_mohafez': Shekh_Mohafez.objects.all(),
        'students': all_students,
    }
    # Check if 'word' is in the request and generate Word document
    if 'word' in request.GET and request.GET['word'] == 'true':
        # Use the last stored values
        name = original_name
        noofguze = original_noofguze
        id_shekh = original_id_shekh
        cs = original_cs
        # Apply filters based on the stored values
        if cs == 'on':
            all_students = all_students.filter(full_name__contains=name)
        else:
            if name==None:
                all_students=all_students
            else:
                all_students = all_students.filter(full_name__icontains=name)
        if id_shekh==None:
            all_students = all_students
        else:
            if int(id_shekh) != 0  :
                all_students = all_students.filter(student_shekh_mohafez__id__icontains=id_shekh)
            else:
                all_students = all_students
        if noofguze == None:
            all_students = all_students
        else:
            if int(noofguze) != 0  :
                all_students = all_students.filter(number_of_Juze__icontains=noofguze)
            else:
                all_students = all_students
        # Update the context with the filtered students
        # Generate Word document here...
        # Return the Word document response
        return generate_word_response(all_students)
    # Store the current values in the session
    request.session['original_name'] = request.GET.get('searchname', None)
    request.session['original_noofguze'] = request.GET.get('no_of_juze_search', None)
    request.session['original_id_shekh'] = request.GET.get('mohafez_select', None)
    request.session['original_cs'] = request.GET.get('cs', 'off')
    return render(request, 'students/search.html', context)
#
def show_results(request):
    return render(request,'students/show_results.html')


